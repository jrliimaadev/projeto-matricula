import os
import random
import sqlite3
import locale
from flask import Flask, render_template,send_file, request, redirect, url_for, session, send_from_directory
from datetime import datetime
from docx import Document
from docx.shared import Pt


app = Flask(__name__)
app.secret_key = 'chave_secreta'
app.config['UPLOAD_FOLDER'] = 'uploads'

os.makedirs('uploads/frente', exist_ok=True)
os.makedirs('uploads/verso', exist_ok=True)

consultores = {
    "Thalita": "https://wa.me/558231420581",
    "Maria": "https://wa.me/558231421590"
}

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/sucesso', methods=['POST'])
def sucesso():
    dados = request.form
    numero_matricula = str(random.randint(100000, 999999))
    data_matricula = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    frente = request.files['documento_frente']
    verso = request.files['documento_verso']

    frente_path = os.path.join('uploads/frente', frente.filename)
    verso_path = os.path.join('uploads/verso', verso.filename)

    frente.save(frente_path)
    verso.save(verso_path)

    conn = sqlite3.connect('banco_dados.db')
    cursor = conn.cursor()
    cursor.execute("""
    INSERT INTO matriculas (
        nome_completo, endereco, cep, cpf, rg, data_nascimento,
        cidade_nascimento, estado_nascimento, tipo_certificado, plano_eja, consultor,
        numero_matricula, documento_frente, documento_verso, data_matricula, status_pagamento
) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
""", (
    dados['nome_completo'], dados['endereco'], dados['cep'],
    dados['cpf'], dados['rg'], dados['data_nascimento'],
    dados['cidade_nascimento'], dados['estado_nascimento'],
    dados['tipo_certificado'], dados['plano_eja'], dados['consultor'],
    numero_matricula, frente.filename, verso.filename, data_matricula, "Não pago"
))

    conn.commit()
    conn.close()

    link_whatsapp = consultores.get(dados['consultor'], "https://wa.me/558231420581")

    return render_template('sucesso.html', numero_matricula=numero_matricula, link_whatsapp=link_whatsapp)

@app.route('/login', methods=['GET', 'POST'])
def login():
    erro = None
    if request.method == 'POST':
        usuario = request.form['usuario']
        senha = request.form['senha']
        conn = sqlite3.connect('banco_dados.db')
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM usuarios WHERE usuario=? AND senha=?", (usuario, senha))
        user = cursor.fetchone()
        conn.close()
        if user:
            session['usuario'] = usuario
            return redirect(url_for('painel'))
        else:
            erro = 'Usuário ou senha inválidos'
    return render_template('login.html', erro=erro)

@app.route('/painel')
def painel():
    if 'usuario' not in session:
        return redirect(url_for('login'))
    conn = sqlite3.connect('banco_dados.db')
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM matriculas")
    matriculas = cursor.fetchall()
    conn.close()
    return render_template('painel.html', matriculas=matriculas, usuario=session['usuario'])

@app.route('/confirmar_pagamento/<int:id>')
def confirmar_pagamento(id):
    if 'usuario' not in session:
        return redirect(url_for('login'))
    conn = sqlite3.connect('banco_dados.db')
    cursor = conn.cursor()
    cursor.execute("UPDATE matriculas SET status_pagamento='Pago' WHERE id=?", (id,))
    conn.commit()
    conn.close()
    return redirect(url_for('painel'))

@app.route('/excluir/<int:id>')
def excluir(id):
    if 'usuario' not in session:
        return redirect(url_for('login'))
    conn = sqlite3.connect('banco_dados.db')
    cursor = conn.cursor()
    cursor.execute("DELETE FROM matriculas WHERE id=?", (id,))
    conn.commit()
    conn.close()
    return redirect(url_for('painel'))

@app.route('/logout')
def logout():
    session.pop('usuario', None)
    return redirect(url_for('login'))

@app.route('/uploads/<tipo>/<nome_arquivo>')
def uploads(tipo, nome_arquivo):
    return send_from_directory(os.path.join('uploads', tipo), nome_arquivo)

from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from docx.shared import Pt

from docx.shared import Pt  # já deve estar no topo do arquivo

def substituir_texto(paragrafo, dados):
    texto_completo = ''.join([run.text for run in paragrafo.runs])
    alterado = False

    for marcador, valor in dados.items():
        if marcador in texto_completo:
            texto_completo = texto_completo.replace(marcador, valor)
            alterado = True

    if alterado:
        for _ in range(len(paragrafo.runs)):
            p_run = paragrafo.runs[0]._element
            p_run.getparent().remove(p_run)

        run = paragrafo.add_run(texto_completo)

        # Aplica formatação especial para o nome do aluno
        if dados.get("{{nome_aluno}}") in texto_completo:
            run.bold = True
            run.font.size = Pt(16)

@app.route('/gerar_declaracao/<int:id>')
def gerar_declaracao(id):
    from docx import Document
    from datetime import datetime
    import os
    import sqlite3

    # Tradução manual dos meses
    meses_pt = {
        'January': 'janeiro', 'February': 'fevereiro', 'March': 'março', 'April': 'abril',
        'May': 'maio', 'June': 'junho', 'July': 'julho', 'August': 'agosto',
        'September': 'setembro', 'October': 'outubro', 'November': 'novembro', 'December': 'dezembro'
    }

    data_atual = datetime.today()
    mes_extenso = meses_pt[data_atual.strftime('%B')]
    data_formatada = f"{data_atual.day} de {mes_extenso} de {data_atual.year}"

    conn = sqlite3.connect('banco_dados.db')
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM matriculas WHERE id=?", (id,))
    aluno = cursor.fetchone()
    print("Ordem dos dados:", aluno)
    # Converter data de nascimento para formato brasileiro
    data_nascimento = datetime.strptime(aluno[6], "%Y-%m-%d").strftime("%d/%m/%Y")

    conn.close()

    if not aluno:
        return "Aluno não encontrado."

    dados = {
        "{{nome_aluno}}": aluno[1].upper(),
        "{{cpf}}": aluno[4],
        "{{rg}}": aluno[5],
        "{{data_nascimento}}": data_nascimento,
        "{{nacionalidade}}": "BRASILEIRO(A)",
        "{{tipo_certificado}}": aluno[9],
        "{{data_hoje}}": data_formatada
    }

    modelo = Document("modelo_declaracao_base.docx")

    for p in modelo.paragraphs:
        substituir_texto(p, dados)

    nome_arquivo = f"declaracao_{aluno[1].replace(' ', '_')}.docx"
    caminho_arquivo = os.path.join("declaracoes", nome_arquivo)

    os.makedirs("declaracoes", exist_ok=True)
    modelo.save(caminho_arquivo)

    return send_file(caminho_arquivo, as_attachment=True)


@app.route('/inserir_valor/<int:id>', methods=['POST'])
def inserir_valor(id):
    valor = request.form.get('valor_pago')
    
    conn = sqlite3.connect('banco_dados.db')
    cursor = conn.cursor()
    cursor.execute("UPDATE matriculas SET valor_pago = ? WHERE id = ?", (valor, id))
    conn.commit()
    conn.close()

    return redirect('/painel')




@app.route('/relatorios')
def relatorios():
    import pandas as pd
    conn = sqlite3.connect('banco_dados.db')
    df = pd.read_sql_query("SELECT consultor, data_matricula, valor_pago FROM matriculas", conn)
    conn.close()

    def extrair_valor(valor):
        try:
            return float(valor.replace("R$", "").replace(".", "").replace(",", ".").strip())
        except:
            return 0.0

    df["valor_pago"] = df["valor_pago"].apply(extrair_valor)
    df["data_matricula"] = pd.to_datetime(df["data_matricula"], errors="coerce")
    df["mes_ano"] = df["data_matricula"].dt.to_period("M")

    relatorio = df.groupby(["mes_ano", "consultor"]).agg(
        total_matriculas=("consultor", "size"),
        valor_total=("valor_pago", "sum")
    ).reset_index()

    relatorio["comissao"] = relatorio["valor_total"] * 0.10

    return render_template("relatorios.html", relatorio=relatorio)


if __name__ == '__main__':
    app.run(debug=True)
